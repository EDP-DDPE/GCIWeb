# app/utils/docx_helper.py
import os
import logging
from io import BytesIO
from docx import Document
from docxtpl import DocxTemplate
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------
# 2️⃣  Funções auxiliares
# ----------------------------------------------------------------------
def remover_content_control_por_tag(doc, tag_name):
    """Remove <w:sdt> cujo <w:tag w:val="tag_name">."""
    removidos = 0
    for sdt in list(doc.element.iter(qn('w:sdt'))):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is None:
            continue
        tag_element = sdtPr.find(qn('w:tag'))
        if tag_element is None:
            continue
        if tag_element.get(qn('w:val')) == tag_name:
            parent = sdt.getparent()
            if parent is not None:
                parent.remove(sdt)
                removidos += 1
    return removidos

def aplicar_fonte_em_todos_runs(doc):
    """
    Mantida para compatibilidade, mas não faz mais nada —
    a fonte FT Base é herdada do estilo do template.
    """
    pass

def ajustar_tamanho_tabela_word(doc, indice_tabela, qtd_linhas):
    
    if indice_tabela >= len(doc.tables):
        return

    tabela = doc.tables[indice_tabela]
    linhas_atuais = len(tabela.rows)
    while linhas_atuais - 1 > qtd_linhas:
        tabela._element.remove(tabela.rows[1]._element)
        linhas_atuais -= 1


def ajustar_tamanho_tabela_alternativas(doc, indice_tabela, qtd_linhas):
    
    if indice_tabela >= len(doc.tables):
        return

    tabela = doc.tables[indice_tabela]
    linhas_atuais = len(tabela.rows)

    mapa = {1: 5, 2: 13, 3: 20, 4: 27}
    linhas_alvo = mapa.get(qtd_linhas, 33)

    while linhas_atuais > linhas_alvo:
        tabela._element.remove(tabela.rows[0]._element)
        linhas_atuais -= 1


def inserir_texto_com_negrito(run_elem, texto_formatado):
    parent_p = run_elem.getparent()
    partes = texto_formatado.split('**')

    # Captura o rPr original do run base (para herdar fonte/tamanho)
    rPr_base = run_elem.find(qn('w:rPr'))

    for i, parte in enumerate(partes):
        if not parte:
            continue
        is_bold = i % 2 != 0

        if i == 0:
            run_to_use = run_elem
        else:
            novo_run = OxmlElement('w:r')
            parent_p.append(novo_run)

            # Copia o rPr base para herdar fonte do template
            if rPr_base is not None:
                from copy import deepcopy
                novo_run.insert(0, deepcopy(rPr_base))

            run_to_use = novo_run

        # Negrito
        if is_bold:
            rPr = run_to_use.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run_to_use.insert(0, rPr)
            b = OxmlElement('w:b')
            rPr.append(b)

        # Texto
        text_node = OxmlElement('w:t')
        text_node.set(qn('xml:space'), 'preserve')
        text_node.text = parte
        run_to_use.append(text_node)

def aplicar_fonte_ft_base(run_elem):
    """
    Não força fonte — o run herda o estilo do parágrafo/documento (FT Base).
    Mantida apenas para não quebrar as chamadas existentes.
    """
    pass


def inserir_imagem_no_content_control(doc, tag_name, caminho_imagem, largura_cm=10):
    """
    Substitui um content‑control do tipo imagem (tag_name) por uma imagem.
    """
    for sdt in list(doc.element.iter(qn('w:sdt'))):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is None:
            continue
        tag = sdtPr.find(qn('w:tag'))
        if tag is None or tag.get(qn('w:val')) != tag_name:
            continue

        sdtContent = sdt.find(qn('w:sdtContent'))
        if sdtContent is not None:
            sdt.remove(sdtContent)

        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        drawing = OxmlElement('w:drawing')

        from docx.shared import Cm
        from docx.oxml import parse_xml

        tmp_doc = Document()
        tmp_doc.add_picture(caminho_imagem, width=Cm(largura_cm))
        tmp_xml = tmp_doc.paragraphs[0]._p.xml
        drawing = parse_xml(tmp_xml)

        r.append(drawing)
        p.append(r)
        sdt.append(p)
        break


# ----------------------------------------------------------------------
# 3️⃣  Processamento de todos os content‑controls (tags) do documento
# ----------------------------------------------------------------------
def processar_elementos(elementos, dados, doc=None):
    """
    Percorre a árvore XML do documento e preenche:
    * textos simples ({{campo}})
    * imagens (dado do tipo dict com {"tipo":"imagem", "caminho":...})
    * remove controles vazios quando necessário
    """
    for element in elementos:
        if element.tag != qn('w:sdt'):
            continue

        tag_name = None
        for child in element:
            if child.tag == qn('w:sdtPr'):
                tag_el = child.find(qn('w:tag'))
                if tag_el is not None:
                    tag_name = tag_el.get(qn('w:val'))
                    break

        if not tag_name:
            continue

        if tag_name in ("estudos_solar", "estudos_hidreletrica", "estudos_termoeletrica"):
            continue

        content = element.find(qn('w:sdtContent'))
        if content is None:
            continue

        # Imagem
        if (doc is not None and tag_name in dados and
                isinstance(dados[tag_name], dict) and
                dados[tag_name].get("tipo") == "imagem"):
            inserir_imagem_no_content_control(
                doc,
                tag_name,
                dados[tag_name]["caminho"],
                largura_cm=dados[tag_name].get("largura_cm", 10),
            )
            continue

        # Texto simples
        if tag_name in dados and not (
                isinstance(dados[tag_name], dict) and
                dados[tag_name].get("tipo") == "imagem"):
            text_value = str(dados[tag_name])

            for t in list(content.iter(qn('w:t'))):
                t.text = ""

            first_run = None
            for t in content.iter(qn('w:t')):
                first_run = t.getparent()
                break
            if first_run is None:
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                r.append(t)
                p.append(r)
                content.append(p)
                first_run = r

            partes = text_value.split('\n')
            parent_p = first_run.getparent()
            for r in list(parent_p):
                if r.tag == qn('w:r') and r is not first_run:
                    parent_p.remove(r)

            for i, parte in enumerate(partes):
                if i > 0:
                    br_run = OxmlElement('w:r')
                    br = OxmlElement('w:br')
                    br_run.append(br)
                    parent_p.append(br_run)

                    novo_run = OxmlElement('w:r')
                    parent_p.append(novo_run)
                    run_para_uso = novo_run
                else:
                    run_para_uso = first_run

                inserir_texto_com_negrito(run_para_uso, parte)


# ----------------------------------------------------------------------
# 4️⃣  Função principal que preenche o template e devolve BytesIO
# ----------------------------------------------------------------------
def preencher_template(template_name: str, context: dict) -> BytesIO:
    """
    Carrega o template .docx, preenche todos os content‑controls (corpo,
    cabeçalhos e rodapés) e devolve o documento em memória.
    """
    #template_path = _load_template(template_name)

    doc = Document(template_name)

    # Tratando documentos com etapa única
    if context["multiplas_etapas"] == False:
        idx = 0
        qtd_linhas_tabela_demanda = 2

        ajustar_tamanho_tabela_word(doc, idx, qtd_linhas_tabela_demanda)

        context["etapa_L5"] = "Atual"
        context["data_L5"] = "Atual"
        context["etapa_L6"] = "Única"
    else:
        del context["carga_ponta_L5"]
        del context["carga_fora_ponta_L5"]
        del context["geracao_ponta_L5"]
        del context["geracao_fora_ponta_L5"]
        del context["data_L6"]
        del context["carga_ponta_L6"]
        del context["carga_fora_ponta_L6"]
        del context["geracao_ponta_L6"]
        del context["geracao_fora_ponta_L6"]

    # Tratando documento com alternativa única
    if context["alternativa_unica"] == True:
        context["das_alternativas"] = "**DA** **ALTERNATIVA**"
        context["das_alternativas_2"] = "da alternativa estudada"
        context["as_alternativas_viaveis"] = "a alternativa viável"
        context["das_alternativas_3"] = "da alternativa avaliada"

        idx = 1
        qtd_linhas_tabela_demanda = 1
        ajustar_tamanho_tabela_word(doc, idx, qtd_linhas_tabela_demanda)

        context["conclusao"] = "Observa-se que pela alternativa estudada o atendimento é viável."

        idx = 2
        qtd_linhas_tabela_demanda = 1
        ajustar_tamanho_tabela_alternativas(doc, idx, qtd_linhas_tabela_demanda)

        context["num_A4"] = "1"

    # Para geração, manter texto de estudos
    if context["c_g"] == "G":
        if context["tipo_geracao"] == "Fotovoltaica":
            remover_content_control_por_tag(doc, "estudos_hidreletrica")
            remover_content_control_por_tag(doc, "estudos_termoeletrica")
            context["tipo_geracao_1"] = "SOLAR"
        elif context["tipo_geracao"] == "Hidrelétrica":
            remover_content_control_por_tag(doc, "estudos_solar")
            remover_content_control_por_tag(doc, "estudos_termoeletrica")
            context["tipo_geracao_1"] = "HIDRELÉTRICA"
        elif context["tipo_geracao"] == "Termoelétrica":
            remover_content_control_por_tag(doc, "estudos_solar")
            remover_content_control_por_tag(doc, "estudos_hidreletrica")
            context["tipo_geracao_1"] = "TERMOELÉTRICA"

    # Tratar caso em que não tem instalação
    if context["instalacao"] == 0 or context["instalacao"] == "-":
        remover_content_control_por_tag(doc, "instalacao")
    else:
        context["instalacao"] = f"(Instalação: {context['instalacao']})"

    processar_elementos(doc.element.iter(), context, doc)

    for section in doc.sections:
        if section.header:
            processar_elementos(section.header._element.iter(), context, doc)
        if section.footer:
            processar_elementos(section.footer._element.iter(), context, doc)

    out_io = BytesIO()
    doc.save(out_io)
    out_io.seek(0)
    return out_io